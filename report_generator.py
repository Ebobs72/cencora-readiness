#!/usr/bin/env python3
"""
Report generator for Launch Readiness assessments.
Generates three types of reports:
1. Baseline Report (PRE assessment only)
2. Progress Report (PRE vs POST comparison)
3. Impact Report (Cohort summary)
All reports use Cencora branding and house style.
"""
import io
import os
import tempfile
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import numpy as np
from framework import (
    INDICATORS, INDICATOR_DESCRIPTIONS, INDICATOR_COLOURS,
    ITEMS, OPEN_QUESTIONS_PRE, OPEN_QUESTIONS_POST, RATING_SCALE,
    FOCUS_TAGS, get_items_for_indicator, get_items_by_focus
)
from theme_extractor import ThemeExtractor, format_themes_for_report, format_insight_themes
# Cencora brand colours (hex for matplotlib, RGB for docx)
COLOURS_HEX = {
    'purple': '#461E96',
    'cyan': '#00B4E6',
    'magenta': '#E6008C',
    'green': '#00DC8C',
    'orange': '#FFA400',
    'dark_grey': '#3B3B3B',
    'mid_grey': '#6E6E6E',
    'light_grey': '#F5F5F5',
    'cream': '#FDF6E3',
    'light_green': '#E8F5E9',
    'success_green': '#007F50',
    'white': '#FFFFFF'
}
COLOURS_RGB = {
    'purple': RGBColor(0x46, 0x1E, 0x96),
    'cyan': RGBColor(0x00, 0xB4, 0xE6),
    'magenta': RGBColor(0xE6, 0x00, 0x8C),
    'green': RGBColor(0x00, 0xDC, 0x8C),
    'orange': RGBColor(0xFF, 0xA4, 0x00),
    'dark_grey': RGBColor(0x3B, 0x3B, 0x3B),
    'mid_grey': RGBColor(0x6E, 0x6E, 0x6E),
    'light_grey': RGBColor(0xF5, 0xF5, 0xF5),
    'cream': RGBColor(0xFD, 0xF6, 0xE3),
    'light_green': RGBColor(0xE8, 0xF5, 0xE9),
    'success_green': RGBColor(0x00, 0x7F, 0x50),
    'white': RGBColor(0xFF, 0xFF, 0xFF)
}
INDICATOR_COLOUR_MAP = {
    'Self-Readiness': 'purple',
    'Practical Readiness': 'cyan',
    'Professional Readiness': 'magenta',
    'Team Readiness': 'green'
}

# Column widths for item detail tables (in twips: 1440 twips = 1 inch)
# 5-col: #, Statement, Focus, Bar, Score  (total ~9000 twips for A4 content area)
COL_WIDTHS_5 = [504, 5040, 1152, 1296, 792]   # 0.35", 3.5", 0.8", 0.9", 0.55"
# 7-col: #, Statement, Focus, Pre, Post, Bar, Change (Progress report)
COL_WIDTHS_7 = [432, 4032, 1296, 576, 576, 1296, 792]  # totals 9000 (Focus now 0.9")
# Logo path - multiple fallback locations for different environments
def get_logo_path():
    """Get the logo path, checking multiple locations for compatibility."""
    # Current file's directory
    current_dir = Path(__file__).parent
    
    # Possible locations (includes Streamlit Cloud mount path)
    candidates = [
        current_dir / 'assets' / 'cencora_logo.png',  # Standard location
        current_dir / 'cencora_logo.png',              # Root of app
        Path('/mount/src/cencora-readiness/assets/cencora_logo.png'),  # Streamlit Cloud
        Path('/mount/src/cencora-readiness/cencora_logo.png'),         # Streamlit Cloud alt
        Path('assets') / 'cencora_logo.png',           # Relative to cwd
        Path('cencora_logo.png'),                      # Current directory
    ]
    
    for path in candidates:
        if path.exists():
            return path
    
    return None
class ReportGenerator:
    def __init__(self, db):
        self.db = db
        self.theme_extractor = ThemeExtractor()
        self.logo_path = get_logo_path()
    
    # =========== CHART GENERATION ===========
    
    def _create_radar_chart(self, scores: dict, output_path: str):
        """Create a 4-axis radar chart matching the approved sample style.
        
        FIX: Uses np.linspace with theta_offset/direction instead of manual
        angle array, which was only rendering one quadrant of the chart.
        """
        
        indicators = list(INDICATORS.keys())
        values = [scores.get(ind, 0) for ind in indicators]
        
        # Chart setup
        fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
        fig.patch.set_facecolor('white')
        
        # Distribute angles evenly around full circle
        angles = np.linspace(0, 2 * np.pi, 4, endpoint=False)
        ax.set_theta_offset(np.pi / 2)   # Self-Readiness at top (12 o'clock)
        ax.set_theta_direction(-1)        # Clockwise order
        
        # Close the polygon
        values_closed = values + [values[0]]
        angles_closed = np.append(angles, angles[0])
        
        # Draw grid circles
        ax.set_ylim(0, 6)
        ax.set_yticks([1, 2, 3, 4, 5, 6])
        ax.set_yticklabels(['1', '2', '3', '4', '5', '6'], fontsize=9, color='#999999')
        ax.set_rlabel_position(45)
        ax.yaxis.grid(True, color='#999999', linewidth=1.2, alpha=0.9)
        ax.xaxis.grid(True, color='#999999', linewidth=1.2, alpha=0.9)
        
        # Draw the data polygon
        ax.fill(angles_closed, values_closed, color=COLOURS_HEX['purple'], alpha=0.15)
        ax.plot(angles_closed, values_closed, color=COLOURS_HEX['purple'], linewidth=3.5)
        
        # Draw data points with indicator colours
        indicator_colours = [COLOURS_HEX['purple'], COLOURS_HEX['cyan'], 
                           COLOURS_HEX['magenta'], COLOURS_HEX['green']]
        for angle, value, colour in zip(angles, values, indicator_colours):
            ax.scatter(angle, value, color=colour, s=200, zorder=5, edgecolors='white', linewidths=2)
        
        # Labels
        ax.set_xticks(angles)
        ax.set_xticklabels([])  # Remove default labels
        
        # Add custom positioned labels
        label_distance = 7.8
        alignments = [
            ('center', 'bottom'),  # Top (Self-Readiness)
            ('left', 'center'),    # Right (Practical Readiness)
            ('center', 'top'),     # Bottom (Professional Readiness)
            ('right', 'center'),   # Left (Team Readiness)
        ]
        for i, (ind, colour) in enumerate(zip(indicators, indicator_colours)):
            ha, va = alignments[i]
            ax.text(angles[i], label_distance, ind, ha=ha, va=va,
                    fontsize=18, fontweight='bold', color=colour)
        
        ax.spines['polar'].set_visible(False)
        
        plt.savefig(output_path, dpi=150, bbox_inches='tight', facecolor='white', 
                   edgecolor='none', pad_inches=0.3)
        plt.close()
    
    def _create_comparison_radar_chart(self, pre_scores: dict, post_scores: dict, output_path: str):
        """Create a comparison radar chart (pre dashed grey, post solid green).
        
        FIX: Same linspace fix as _create_radar_chart.
        """
        
        indicators = list(INDICATORS.keys())
        pre_values = [pre_scores.get(ind, 0) for ind in indicators]
        post_values = [post_scores.get(ind, 0) for ind in indicators]
        
        fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
        fig.patch.set_facecolor('white')
        
        # Distribute angles evenly around full circle
        angles = np.linspace(0, 2 * np.pi, 4, endpoint=False)
        ax.set_theta_offset(np.pi / 2)   # Self-Readiness at top
        ax.set_theta_direction(-1)        # Clockwise
        
        pre_closed = pre_values + [pre_values[0]]
        post_closed = post_values + [post_values[0]]
        angles_closed = np.append(angles, angles[0])
        
        ax.set_ylim(0, 6)
        ax.set_yticks([1, 2, 3, 4, 5, 6])
        ax.set_yticklabels(['1', '2', '3', '4', '5', '6'], fontsize=9, color='#999999')
        ax.set_rlabel_position(45)
        ax.yaxis.grid(True, color='#999999', linewidth=1.2, alpha=0.9)
        ax.xaxis.grid(True, color='#999999', linewidth=1.2, alpha=0.9)
        
        # PRE polygon (dashed, grey)
        ax.fill(angles_closed, pre_closed, color='#999999', alpha=0.1)
        ax.plot(angles_closed, pre_closed, color='#999999', linewidth=2.5, linestyle='--')
        
        # POST polygon (solid, green)
        ax.fill(angles_closed, post_closed, color=COLOURS_HEX['green'], alpha=0.15)
        ax.plot(angles_closed, post_closed, color=COLOURS_HEX['green'], linewidth=3.5)
        
        # Points
        indicator_colours = [COLOURS_HEX['purple'], COLOURS_HEX['cyan'],
                           COLOURS_HEX['magenta'], COLOURS_HEX['green']]
        
        for angle, pre_val, post_val, colour in zip(angles, pre_values, post_values, indicator_colours):
            # PRE point (smaller, grey)
            ax.scatter(angle, pre_val, color='#999999', s=80, zorder=4, edgecolors='white', linewidths=1)
            # POST point (larger, coloured)
            ax.scatter(angle, post_val, color=colour, s=200, zorder=5, edgecolors='white', linewidths=2)
        
        ax.set_xticks(angles)
        ax.set_xticklabels([])
        
        # Custom positioned labels
        label_distance = 7.8
        alignments = [
            ('center', 'bottom'),
            ('left', 'center'),
            ('center', 'top'),
            ('right', 'center'),
        ]
        for i, (ind, colour) in enumerate(zip(indicators, indicator_colours)):
            ha, va = alignments[i]
            ax.text(angles[i], label_distance, ind, ha=ha, va=va,
                    fontsize=18, fontweight='bold', color=colour)
        
        # Legend
        from matplotlib.lines import Line2D
        legend_elements = [
            Line2D([0], [0], color='#999999', linestyle='--', linewidth=2, label='Pre-Programme'),
            Line2D([0], [0], color=COLOURS_HEX['green'], linewidth=2.5, label='Post-Programme')
        ]
        ax.legend(handles=legend_elements, loc='lower left', bbox_to_anchor=(-0.1, -0.15), 
                 fontsize=9, frameon=False)
        
        ax.spines['polar'].set_visible(False)
        
        plt.savefig(output_path, dpi=150, bbox_inches='tight', facecolor='white',
                   edgecolor='none', pad_inches=0.3)
        plt.close()
    
    def _create_bar_chart(self, score: float, colour_hex: str, output_path: str, max_score: int = 6):
        """Create a simple horizontal bar for a single score."""
        
        fig, ax = plt.subplots(figsize=(1.5, 0.25))
        fig.patch.set_facecolor('white')
        
        # Background bar
        ax.barh(0, max_score, color='#E8E8E8', height=0.8)
        # Score bar
        ax.barh(0, score, color=colour_hex, height=0.8)
        
        ax.set_xlim(0, max_score)
        ax.set_ylim(-0.5, 0.5)
        ax.axis('off')
        
        plt.savefig(output_path, dpi=100, bbox_inches='tight', facecolor='white',
                   edgecolor='none', pad_inches=0.02)
        plt.close()
    
    def _create_comparison_bar_chart(self, pre_score: float, post_score: float, 
                                     colour_hex: str, output_path: str, max_score: int = 6):
        """Create a stacked comparison bar (pre grey on top, post coloured below)."""
        
        fig, ax = plt.subplots(figsize=(1.5, 0.5))
        fig.patch.set_facecolor('white')
        
        bar_height = 0.35
        
        # PRE bar (top) - light grey
        ax.barh(0.22, max_score, color='#E8E8E8', height=bar_height)
        ax.barh(0.22, pre_score, color='#B0B0B0', height=bar_height)
        
        # POST bar (bottom) - indicator colour
        ax.barh(-0.22, max_score, color='#E8E8E8', height=bar_height)
        ax.barh(-0.22, post_score, color=colour_hex, height=bar_height)
        
        ax.set_xlim(0, max_score)
        ax.set_ylim(-0.55, 0.55)
        ax.axis('off')
        
        plt.savefig(output_path, dpi=100, bbox_inches='tight', facecolor='white',
                   edgecolor='none', pad_inches=0.02)
        plt.close()
    
    # =========== DOCUMENT HELPERS ===========
    
    def _set_cell_shading(self, cell, colour_hex: str):
        """Set cell background colour."""
        colour_hex = colour_hex.replace('#', '')
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), colour_hex)
        cell._tc.get_or_add_tcPr().append(shading)
    
    def _set_cell_margins(self, cell, top=60, bottom=60, left=100, right=100):
        """Set cell margins in twips."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{side}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)
    
    def _add_logo_header(self, doc):
        """Add Cencora logo to document header and page numbers to footer on all pages."""
        section = doc.sections[0]
        
        # Logo in header
        if self.logo_path and self.logo_path.exists():
            try:
                header = section.header
                header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = header_para.add_run()
                # Logo width of 1.5 inches for good visibility
                run.add_picture(str(self.logo_path), width=Inches(1.5))
            except Exception:
                pass  # Skip logo if image is unreadable
        
        # Page numbers in footer: "Page X of Y  |  The Development Catalyst  |  Confidential"
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # "Page "
        run = footer_para.add_run("Page ")
        run.font.size = Pt(8)
        run.font.color.rgb = COLOURS_RGB['mid_grey']
        run.font.name = 'Arial'
        
        # Current page number field
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        run1 = footer_para.add_run()
        run1._r.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run2 = footer_para.add_run()
        run2._r.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run3 = footer_para.add_run()
        run3._r.append(fldChar2)
        
        # " of "
        run = footer_para.add_run(" of ")
        run.font.size = Pt(8)
        run.font.color.rgb = COLOURS_RGB['mid_grey']
        run.font.name = 'Arial'
        
        # Total pages field
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'begin')
        run4 = footer_para.add_run()
        run4._r.append(fldChar3)
        
        instrText2 = OxmlElement('w:instrText')
        instrText2.set(qn('xml:space'), 'preserve')
        instrText2.text = ' NUMPAGES '
        run5 = footer_para.add_run()
        run5._r.append(instrText2)
        
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(qn('w:fldCharType'), 'end')
        run6 = footer_para.add_run()
        run6._r.append(fldChar4)
        
        # "  |  The Development Catalyst  |  Confidential"
        run = footer_para.add_run("  |  The Development Catalyst  |  Confidential")
        run.font.size = Pt(8)
        run.font.color.rgb = COLOURS_RGB['mid_grey']
        run.font.name = 'Arial'
    
    def _create_styled_table(self, doc, headers: list, header_colour_hex: str = '461E96',
                             col_widths: list = None):
        """Create a table with styled header row and optional fixed column widths."""
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Set fixed layout to enforce column widths
        if col_widths:
            tbl = table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'fixed')
            tblPr.append(tblLayout)
        
        # Style header row
        header_row = table.rows[0]
        for i, header_text in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header_text
            self._set_cell_shading(cell, header_colour_hex)
            self._set_cell_margins(cell)
            
            # Set column width
            if col_widths and i < len(col_widths):
                cell.width = col_widths[i]
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(col_widths[i]))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)
            
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.font.bold = True
                    run.font.color.rgb = COLOURS_RGB['white']
                    run.font.size = Pt(9)
                    run.font.name = 'Arial'
        
        return table
    
    def _add_table_row(self, table, values: list, row_index: int, 
                       alignments: list = None, bar_image_path: str = None, bar_col: int = None,
                       col_widths: list = None):
        """Add a data row with alternating colours and optional bar chart."""
        row = table.add_row()
        bg_colour = 'FFFFFF' if row_index % 2 == 0 else 'FDF6E3'
        
        for i, value in enumerate(values):
            cell = row.cells[i]
            
            # Enforce column width on data rows too
            if col_widths and i < len(col_widths):
                cell.width = col_widths[i]
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(col_widths[i]))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)
            
            # Check if this cell should have a bar image
            if bar_image_path and i == bar_col:
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(bar_image_path, width=Inches(0.8))
            else:
                cell.text = str(value) if value is not None else ''
                for para in cell.paragraphs:
                    if alignments and i < len(alignments):
                        para.alignment = alignments[i]
                    for run in para.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Arial'
            
            self._set_cell_shading(cell, bg_colour)
            self._set_cell_margins(cell)
        
        return row
    
    # =========== CALCULATION HELPERS ===========
    
    def _calculate_indicator_scores(self, ratings: dict) -> dict:
        """Calculate average score per indicator from ratings."""
        scores = {}
        for indicator, (start, end) in INDICATORS.items():
            item_scores = [ratings.get(i, 0) for i in range(start, end + 1) if i in ratings]
            if item_scores:
                scores[indicator] = sum(item_scores) / len(item_scores)
            else:
                scores[indicator] = 0
        return scores
    
    def _calculate_overall_score(self, ratings: dict) -> float:
        """Calculate overall average from all 32 ratings."""
        if not ratings:
            return 0
        valid_ratings = [ratings.get(i, 0) for i in range(1, 33) if i in ratings]
        return sum(valid_ratings) / len(valid_ratings) if valid_ratings else 0
    
    def _calculate_focus_scores(self, ratings: dict) -> dict:
        """Calculate average scores by focus area (K/A/C/B)."""
        scores = {}
        for focus in FOCUS_TAGS.keys():
            items = get_items_by_focus(focus)
            item_scores = [ratings.get(i, 0) for i in items if i in ratings]
            if item_scores:
                scores[focus] = sum(item_scores) / len(item_scores)
            else:
                scores[focus] = 0
        return scores
    
    # =========== APPENDIX ===========
    
    def _add_appendix(self, doc):
        """Add appendix page with rating scale definitions and focus area descriptions."""
        doc.add_page_break()
        
        heading = doc.add_paragraph()
        run = heading.add_run("Appendix")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = COLOURS_RGB['purple']
        
        # Rating scale
        sub = doc.add_paragraph()
        run = sub.add_run("Rating Scale")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = COLOURS_RGB['purple']
        
        scale_intro = doc.add_paragraph()
        run = scale_intro.add_run("All statements are rated on a 6-point agreement scale:")
        run.font.size = Pt(9)
        run.font.color.rgb = COLOURS_RGB['mid_grey']
        
        doc.add_paragraph()
        scale_table = self._create_styled_table(doc, ["Score", "Label", "Interpretation"])
        
        scale_data = [
            ("1", "Strongly Disagree", "This is not true for me at all — a clear development priority"),
            ("2", "Disagree", "I don't feel this applies to me yet — needs focused attention"),
            ("3", "Slightly Disagree", "I'm starting to develop here but not yet confident"),
            ("4", "Slightly Agree", "I'm building capability — some evidence but room to grow"),
            ("5", "Agree", "This is generally true for me — a developing strength"),
            ("6", "Strongly Agree", "This is consistently true for me — a clear strength"),
        ]
        
        scale_col_widths = [720, 1800, 6480]
        # Apply widths to header
        for i, w in enumerate(scale_col_widths):
            cell = scale_table.rows[0].cells[i]
            cell.width = w
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(w))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
        
        for idx, (score, label, interp) in enumerate(scale_data):
            self._add_table_row(
                scale_table, [score, label, interp], idx,
                [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
                col_widths=scale_col_widths
            )
        
        # Focus areas
        doc.add_paragraph()
        doc.add_paragraph()
        sub = doc.add_paragraph()
        run = sub.add_run("Focus Areas (BACK)")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = COLOURS_RGB['purple']
        
        focus_intro = doc.add_paragraph()
        run = focus_intro.add_run("Each of the 32 statements measures one of four focus areas. "
                                  "Together they give a rounded picture of readiness:")
        run.font.size = Pt(9)
        run.font.color.rgb = COLOURS_RGB['mid_grey']
        
        doc.add_paragraph()
        focus_table = self._create_styled_table(doc, ["Focus", "What It Measures", "Example"])
        
        focus_data = [
            ("Behaviour", 
             "Observable actions, habits and practices — what you actually do day to day",
             "\"I delegate tasks appropriately rather than taking on too much myself\""),
            ("Awareness", 
             "Recognition of your own patterns, triggers, impact on others and development needs",
             "\"I recognise how my behaviour changes when I am under pressure\""),
            ("Confidence", 
             "Self-belief, comfort in capability, and willingness to step into challenging situations",
             "\"I feel equipped to handle common people management situations\""),
            ("Knowledge", 
             "Understanding of concepts, processes, frameworks and how to apply them",
             "\"I understand how to match my delegation approach to the individual and the task\""),
        ]
        
        focus_col_widths = [1080, 3960, 3960]
        for i, w in enumerate(focus_col_widths):
            cell = focus_table.rows[0].cells[i]
            cell.width = w
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(w))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
        
        for idx, (focus, desc, example) in enumerate(focus_data):
            self._add_table_row(
                focus_table, [focus, desc, example], idx,
                [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT],
                col_widths=focus_col_widths
            )
        
        # Closing note
        doc.add_paragraph()
        note = doc.add_paragraph()
        run = note.add_run("The Readiness Framework was developed by The Development Catalyst "
                          "to measure launch readiness across the dimensions that matter most "
                          "for new leaders and managers.")
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = COLOURS_RGB['mid_grey']
    
    # =========== BASELINE REPORT ===========
    
    def generate_baseline_report(self, participant_id: int) -> io.BytesIO:
        """Generate a Baseline report (PRE assessment only)."""
        
        # Get data
        data = self.db.get_participant_data(participant_id)
        if not data or not data['pre']:
            raise ValueError("No PRE assessment data found for this participant")
        
        participant = data['participant']
        cohort = data['cohort']
        pre_ratings = data['pre']['ratings']
        pre_responses = data['pre']['open_responses']
        pre_date = data['pre']['assessment'].get('completed_at', '')[:10]
        
        # Calculate scores
        indicator_scores = self._calculate_indicator_scores(pre_ratings)
        overall_score = self._calculate_overall_score(pre_ratings)
        
        # Create document
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)
        
        # Add logo header
        self._add_logo_header(doc)
        
        # ===== COVER PAGE =====
        # Add vertical spacing to centre content
        for _ in range(6):
            doc.add_paragraph()
        
        # Programme title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("THE READINESS FRAMEWORK")
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = COLOURS_RGB['purple']
        run.font.name = 'Arial'
        
        # Report type
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Readiness Baseline")
        run.font.size = Pt(18)
        run.font.color.rgb = COLOURS_RGB['magenta']
        run.font.name = 'Arial'
        
        # Decorative line
        doc.add_paragraph()
        line = doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = line.add_run("\u2501" * 40)
        run.font.color.rgb = COLOURS_RGB['purple']
        run.font.size = Pt(10)
        
        doc.add_paragraph()
        
        # Participant name - large and centred
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_para.add_run(participant['name'])
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = COLOURS_RGB['dark_grey']
        run.font.name = 'Arial'
        
        # Role
        role_para = doc.add_paragraph()
        role_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = role_para.add_run(participant.get('role', ''))
        run.font.size = Pt(14)
        run.font.color.rgb = COLOURS_RGB['mid_grey']
        run.font.name = 'Arial'
        
        doc.add_paragraph()
        
        # Cohort and date
        details_para = doc.add_paragraph()
        details_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = details_para.add_run(f"{cohort['name']}  |  Assessment: {pre_date}")
        run.font.size = Pt(11)
        run.font.color.rgb = COLOURS_RGB['mid_grey']
        run.font.name = 'Arial'
        
        gen_para = doc.add_paragraph()
        gen_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = gen_para.add_run(f"Report generated: {datetime.now().strftime('%d %B %Y')}")
        run.font.size = Pt(10)
        run.font.color.rgb = COLOURS_RGB['mid_grey']
        run.font.name = 'Arial'
        
        # Confidentiality note at bottom
        for _ in range(4):
            doc.add_paragraph()
        conf_para = doc.add_paragraph()
        conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = conf_para.add_run("CONFIDENTIAL")
        run.font.size = Pt(9)
        run.font.color.rgb = COLOURS_RGB['mid_grey']
        run.font.name = 'Arial'
        run.bold = True
        
        # ===== PAGE 2: YOUR STARTING POINT =====
        doc.add_page_break()
        
        heading = doc.add_paragraph()
        run = heading.add_run("Your Starting Point")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = COLOURS_RGB['purple']
        
        intro = doc.add_paragraph()
        intro.add_run(f"Welcome to the Launch Readiness programme, {participant['name'].split()[0]}. "
                      f"This report captures your self-assessment before the programme begins. "
                      f"There are no right or wrong answers; this is simply a snapshot of where you see yourself today.")
        
        doc.add_paragraph()
        intro2 = doc.add_paragraph()
        intro2.add_run("The assessment measures your readiness across four key indicators, each containing "
                       "statements about different aspects of your role. Your responses are summarised "
                       "in the radar chart below, followed by a detailed breakdown of each indicator.")
        
        # Radar chart heading
        doc.add_paragraph()
        heading = doc.add_paragraph()
        run = heading.add_run("Your Readiness Profile")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = COLOURS_RGB['purple']
        
        # Radar chart - centered
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            self._create_radar_chart(indicator_scores, tmp.name)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(tmp.name, width=Inches(4.5))
        
        # Scale note
        scale_para = doc.add_paragraph()
        scale_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = scale_para.add_run("Scale: 1-6 (1=Strongly Disagree, 6=Strongly Agree)")
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = COLOURS_RGB['mid_grey']
        
        # Summary table with explicit widths
        doc.add_paragraph()
        summary_col_widths = [5760, 2880]  # 4", 2" - matching proportions
        summary_table = self._create_styled_table(doc, ["Indicator", "Score"],
                                                   col_widths=summary_col_widths)
        
        for i, (ind, score) in enumerate(indicator_scores.items()):
            self._add_table_row(summary_table, [ind, f"{score:.1f}"], i,
                               [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER],
                               col_widths=summary_col_widths)
        
        # Overall row - centred score
        overall_row = summary_table.add_row()
        overall_row.cells[0].text = "OVERALL"
        overall_row.cells[1].text = f"{overall_score:.1f}"
        for j, cell in enumerate(overall_row.cells):
            self._set_cell_shading(cell, 'F5F5F5')
            self._set_cell_margins(cell)
            # Set width
            if j < len(summary_col_widths):
                cell.width = summary_col_widths[j]
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(summary_col_widths[j]))
                tcW.set(qn('w:type'), 'dxa')
                tcPr.append(tcW)
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 1 else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(9)
                    run.font.name = 'Arial'
        
        # Score interpretation key
        doc.add_paragraph()
        key_para = doc.add_paragraph()
        key_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = key_para.add_run("Reading your scores:  ")
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = COLOURS_RGB['dark_grey']
        run.font.name = 'Arial'
        run = key_para.add_run("1\u20132 = Development Priority    3\u20134 = Building    5\u20136 = Strength")
        run.font.size = Pt(8)
        run.font.color.rgb = COLOURS_RGB['mid_grey']
        run.font.name = 'Arial'
        for indicator, (start, end) in INDICATORS.items():
            colour_name = INDICATOR_COLOUR_MAP.get(indicator, 'purple')
            colour_hex = COLOURS_HEX[colour_name]
            
            # Page break before each indicator to prevent table splitting
            doc.add_page_break()
            
            # Indicator heading
            heading = doc.add_paragraph()
            run = heading.add_run(indicator)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = COLOURS_RGB[colour_name]
            
            # Description and average
            desc = doc.add_paragraph()
            run = desc.add_run(INDICATOR_DESCRIPTIONS.get(indicator, ''))
            run.italic = True
            run.font.color.rgb = COLOURS_RGB['mid_grey']
            run.font.size = Pt(9)
            desc.add_run("  |  ")
            desc.add_run("Dimension Average: ")
            run = desc.add_run(f"{indicator_scores.get(indicator, 0):.1f}")
            run.bold = True
            
            # Items table with bar charts
            items_table = self._create_styled_table(
                doc, ["#", "Statement", "Focus", "", "Score"],
                colour_hex.replace('#', ''),
                col_widths=COL_WIDTHS_5
            )
            
            for idx, item_num in enumerate(range(start, end + 1)):
                item = ITEMS[item_num]
                score = pre_ratings.get(item_num, 0)
                
                # Create bar chart for this item
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                    self._create_bar_chart(score, colour_hex, tmp.name)
                    self._add_table_row(
                        items_table,
                        [str(item_num), item['text'], item['focus'], None, str(score)],
                        idx,
                        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                         WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                         WD_ALIGN_PARAGRAPH.CENTER],
                        bar_image_path=tmp.name,
                        bar_col=3,
                        col_widths=COL_WIDTHS_5
                    )
            
            doc.add_paragraph()
        
        # Overall Readiness items
        heading = doc.add_paragraph()
        run = heading.add_run("Overall Readiness")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = COLOURS_RGB['purple']
        
        overall_table = self._create_styled_table(doc, ["#", "Statement", "Focus", "", "Score"],
                                                    col_widths=COL_WIDTHS_5)
        
        for idx, item_num in enumerate([31, 32]):
            item = ITEMS[item_num]
            score = pre_ratings.get(item_num, 0)
            
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                self._create_bar_chart(score, COLOURS_HEX['purple'], tmp.name)
                self._add_table_row(
                    overall_table,
                    [str(item_num), item['text'], item['focus'], None, str(score)],
                    idx,
                    [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                     WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                     WD_ALIGN_PARAGRAPH.CENTER],
                    bar_image_path=tmp.name,
                    bar_col=3,
                    col_widths=COL_WIDTHS_5
                )
        
        # Page break before reflections
        doc.add_page_break()
        
        # Open responses
        heading = doc.add_paragraph()
        run = heading.add_run("Your Reflections")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = COLOURS_RGB['purple']
        
        intro = doc.add_paragraph()
        run = intro.add_run("Your responses to the open questions before the programme:")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = COLOURS_RGB['mid_grey']
        
        for q_num, question in OPEN_QUESTIONS_PRE.items():
            # Question in a shaded box
            q_table = doc.add_table(rows=1, cols=1)
            q_table.style = 'Table Grid'
            q_cell = q_table.rows[0].cells[0]
            q_cell.text = question
            self._set_cell_shading(q_cell, 'F5F5F5')
            self._set_cell_margins(q_cell, 80, 80, 120, 120)
            for para in q_cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
                    run.font.color.rgb = COLOURS_RGB['purple']
            
            # Response
            response = pre_responses.get(q_num, "No response provided")
            r_para = doc.add_paragraph()
            run = r_para.add_run(response)
            run.italic = True
            run.font.size = Pt(10)
            run.font.name = 'Arial'
            r_para.paragraph_format.space_after = Pt(12)
        
        # Closing note
        doc.add_paragraph()
        closing = doc.add_paragraph()
        run = closing.add_run("Keep this report - you'll revisit it after the programme to see how far you've come.")
        run.italic = True
        run.font.color.rgb = COLOURS_RGB['mid_grey']
        run.font.size = Pt(9)
        
        # Appendix
        self._add_appendix(doc)
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    # =========== PROGRESS REPORT ===========
    
    def generate_progress_report(self, participant_id: int, cohort_id: int) -> io.BytesIO:
        """Generate a Progress report (PRE vs POST with cohort comparison)."""
        
        # Get participant data
        data = self.db.get_participant_data(participant_id)
        if not data or not data['pre'] or not data['post']:
            raise ValueError("Both PRE and POST assessments required for Progress report")
        
        participant = data['participant']
        cohort = data['cohort']
        pre_ratings = data['pre']['ratings']
        post_ratings = data['post']['ratings']
        pre_responses = data['pre']['open_responses']
        post_responses = data['post']['open_responses']
        pre_date = data['pre']['assessment'].get('completed_at', '')[:10]
        post_date = data['post']['assessment'].get('completed_at', '')[:10]
        
        # Calculate scores
        pre_indicator_scores = self._calculate_indicator_scores(pre_ratings)
        post_indicator_scores = self._calculate_indicator_scores(post_ratings)
        pre_overall = self._calculate_overall_score(pre_ratings)
        post_overall = self._calculate_overall_score(post_ratings)
        
        # Get cohort averages
        cohort_avgs = self.db.get_cohort_averages(cohort_id, 'POST')
        cohort_indicator_scores = {}
        for indicator, (start, end) in INDICATORS.items():
            item_avgs = [cohort_avgs.get(i, {}).get('avg', 0) for i in range(start, end + 1)]
            cohort_indicator_scores[indicator] = sum(item_avgs) / len(item_avgs) if item_avgs else 0
        cohort_overall = sum(cohort_avgs.get(i, {}).get('avg', 0) for i in range(1, 33)) / 32
        
        # Create document
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)
        
        self._add_logo_header(doc)
        
        # ===== COVER PAGE =====
        for _ in range(6):
            doc.add_paragraph()
        
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("THE READINESS FRAMEWORK")
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = COLOURS_RGB['purple']
        run.font.name = 'Arial'
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run("Readiness Progress")
        run.font.size = Pt(18)
        run.font.color.rgb = COLOURS_RGB['magenta']
        run.font.name = 'Arial'
        
        doc.add_paragraph()
        line = doc.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = line.add_run("\u2501" * 40)
        run.font.color.rgb = COLOURS_RGB['purple']
        run.font.size = Pt(10)
        
        doc.add_paragraph()
        
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_para.add_run(participant['name'])
        run.bold = True
        run.font.size = Pt(22)
        run.font.color.rgb = COLOURS_RGB['dark_grey']
        run.font.name = 'Arial'
        
        role_para = doc.add_paragraph()
        role_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = role_para.add_run(participant.get('role', ''))
        run.font.size = Pt(14)
        run.font.color.rgb = COLOURS_RGB['mid_grey']
        run.font.name = 'Arial'
        
        doc.add_paragraph()
        
        details_para = doc.add_paragraph()
        details_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = details_para.add_run(f"{cohort['name']}  |  Pre: {pre_date}  |  Post: {post_date}")
        run.font.size = Pt(11)
        run.font.color.rgb = COLOURS_RGB['mid_grey']
        run.font.name = 'Arial'
        
        gen_para = doc.add_paragraph()
        gen_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = gen_para.add_run(f"Report generated: {datetime.now().strftime('%d %B %Y')}")
        run.font.size = Pt(10)
        run.font.color.rgb = COLOURS_RGB['mid_grey']
        run.font.name = 'Arial'
        
        for _ in range(4):
            doc.add_paragraph()
        conf_para = doc.add_paragraph()
        conf_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = conf_para.add_run("CONFIDENTIAL")
        run.font.size = Pt(9)
        run.font.color.rgb = COLOURS_RGB['mid_grey']
        run.font.name = 'Arial'
        run.bold = True
        
        # ===== PAGE 2: YOUR PROGRESS =====
        doc.add_page_break()
        
        heading = doc.add_paragraph()
        run = heading.add_run("Your Progress")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = COLOURS_RGB['purple']
        
        change = post_overall - pre_overall
        intro = doc.add_paragraph()
        intro.add_run(f"Congratulations, {participant['name'].split()[0]}! Your assessment shows meaningful "
                      f"growth across all four Readiness Indicators. Your overall readiness improved from "
                      f"{pre_overall:.1f} to {post_overall:.1f} (+{change:.1f}).")
        
        doc.add_paragraph()
        intro2 = doc.add_paragraph()
        intro2.add_run("The comparison radar chart below shows your pre-programme profile (dashed grey) "
                       "alongside your post-programme profile (solid green), followed by detailed "
                       "breakdowns for each indicator.")
        
        # Comparison radar chart
        doc.add_paragraph()
        heading = doc.add_paragraph()
        run = heading.add_run("Your Growth Profile")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = COLOURS_RGB['purple']
        
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            self._create_comparison_radar_chart(pre_indicator_scores, post_indicator_scores, tmp.name)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(tmp.name, width=Inches(4.5))
        
        # Summary table
        doc.add_paragraph()
        summary_table = self._create_styled_table(doc, ["Indicator", "Pre", "Post", "Change", "Cohort"])
        
        for i, indicator in enumerate(INDICATORS.keys()):
            pre = pre_indicator_scores.get(indicator, 0)
            post = post_indicator_scores.get(indicator, 0)
            change = post - pre
            cohort_avg = cohort_indicator_scores.get(indicator, 0)
            change_str = f"+{change:.1f}" if change > 0 else f"{change:.1f}"
            
            self._add_table_row(summary_table,
                               [indicator, f"{pre:.1f}", f"{post:.1f}", change_str, f"{cohort_avg:.1f}"],
                               i,
                               [WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 4)
        
        # Overall row
        change = post_overall - pre_overall
        change_str = f"+{change:.1f}" if change > 0 else f"{change:.1f}"
        overall_row = summary_table.add_row()
        for j, val in enumerate(["OVERALL", f"{pre_overall:.1f}", f"{post_overall:.1f}", change_str, f"{cohort_overall:.1f}"]):
            overall_row.cells[j].text = val
            self._set_cell_shading(overall_row.cells[j], 'F5F5F5')
            self._set_cell_margins(overall_row.cells[j])
            for para in overall_row.cells[j].paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 else WD_ALIGN_PARAGRAPH.LEFT
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(9)
        
        # Cohort note
        note = doc.add_paragraph()
        run = note.add_run("Cohort = Average of all participants  |  Bar: grey = Pre, coloured = Post")
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = COLOURS_RGB['mid_grey']
        
        # Score interpretation key
        key_para = doc.add_paragraph()
        key_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = key_para.add_run("Reading your scores:  ")
        run.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = COLOURS_RGB['dark_grey']
        run.font.name = 'Arial'
        run = key_para.add_run("1\u20132 = Development Priority    3\u20134 = Building    5\u20136 = Strength")
        run.font.size = Pt(8)
        run.font.color.rgb = COLOURS_RGB['mid_grey']
        run.font.name = 'Arial'
        
        # ===== GROWTH HIGHLIGHTS =====
        doc.add_page_break()
        
        heading = doc.add_paragraph()
        run = heading.add_run("Your Growth at a Glance")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = COLOURS_RGB['purple']
        
        intro = doc.add_paragraph()
        intro.add_run("Before diving into the detail, here is a summary of where you showed the "
                      "biggest shifts and where there is still room to grow.")
        
        # Calculate item-level changes and sort
        item_changes = []
        for item_num in range(1, 33):
            pre_score = pre_ratings.get(item_num, 0)
            post_score = post_ratings.get(item_num, 0)
            item_change = post_score - pre_score
            item = ITEMS[item_num]
            item_changes.append({
                'num': item_num,
                'text': item['text'],
                'focus': item['focus'],
                'pre': pre_score,
                'post': post_score,
                'change': item_change
            })
        
        # Sort by change descending for biggest growth
        sorted_by_growth = sorted(item_changes, key=lambda x: x['change'], reverse=True)
        # Sort by post score ascending for areas to develop
        sorted_by_post = sorted(item_changes, key=lambda x: x['post'])
        
        # Biggest Growth (top 5 items by change)
        doc.add_paragraph()
        sub = doc.add_paragraph()
        run = sub.add_run("\u2B06  Your Biggest Growth Areas")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = COLOURS_RGB['success_green']
        
        growth_table = self._create_styled_table(
            doc, ["#", "Statement", "Pre", "Post", "Change"],
            '007F50'  # Cencora dark green
        )
        growth_col_widths = [432, 5760, 864, 864, 1080]
        # Apply widths to header
        for i, w in enumerate(growth_col_widths):
            cell = growth_table.rows[0].cells[i]
            cell.width = w
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(w))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
        
        for idx, item in enumerate(sorted_by_growth[:5]):
            change_str = f"+{item['change']}" if item['change'] > 0 else str(item['change'])
            self._add_table_row(
                growth_table,
                [str(item['num']), item['text'], str(item['pre']), str(item['post']), change_str],
                idx,
                [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                 WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                 WD_ALIGN_PARAGRAPH.CENTER],
                col_widths=growth_col_widths
            )
        
        # Areas still to develop (lowest 5 post scores)
        doc.add_paragraph()
        sub = doc.add_paragraph()
        run = sub.add_run("\u27A1  Areas for Continued Development")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = COLOURS_RGB['orange']
        
        develop_intro = doc.add_paragraph()
        run = develop_intro.add_run("These are the items where your post-programme scores were lowest. "
                                    "They represent your best opportunities for continued growth.")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = COLOURS_RGB['mid_grey']
        
        develop_table = self._create_styled_table(
            doc, ["#", "Statement", "Pre", "Post", "Change"],
            'FFA400'  # orange
        )
        # Apply widths to header
        for i, w in enumerate(growth_col_widths):
            cell = develop_table.rows[0].cells[i]
            cell.width = w
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(w))
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
        
        for idx, item in enumerate(sorted_by_post[:5]):
            change_str = f"+{item['change']}" if item['change'] > 0 else str(item['change'])
            self._add_table_row(
                develop_table,
                [str(item['num']), item['text'], str(item['pre']), str(item['post']), change_str],
                idx,
                [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                 WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                 WD_ALIGN_PARAGRAPH.CENTER],
                col_widths=growth_col_widths
            )
        
        # ===== DETAILED COMPARISON BY INDICATOR =====
        # (Each indicator gets its own page)
        
        for indicator, (start, end) in INDICATORS.items():
            colour_name = INDICATOR_COLOUR_MAP.get(indicator, 'purple')
            colour_hex = COLOURS_HEX[colour_name]
            
            pre_avg = pre_indicator_scores.get(indicator, 0)
            post_avg = post_indicator_scores.get(indicator, 0)
            change = post_avg - pre_avg
            
            # Page break before each indicator to prevent table splitting
            doc.add_page_break()
            
            heading = doc.add_paragraph()
            run = heading.add_run(indicator)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = COLOURS_RGB[colour_name]
            
            desc = doc.add_paragraph()
            run = desc.add_run(INDICATOR_DESCRIPTIONS.get(indicator, ''))
            run.italic = True
            run.font.color.rgb = COLOURS_RGB['mid_grey']
            run.font.size = Pt(9)
            desc.add_run(f"  |  Pre: {pre_avg:.1f} -> Post: {post_avg:.1f} ")
            change_run = desc.add_run(f"(+{change:.1f})" if change > 0 else f"({change:.1f})")
            change_run.bold = True
            if change > 0:
                change_run.font.color.rgb = COLOURS_RGB['success_green']
            
            # Items table
            items_table = self._create_styled_table(
                doc, ["#", "Statement", "Focus", "Pre", "Post", "", "Change"],
                colour_hex.replace('#', ''),
                col_widths=COL_WIDTHS_7
            )
            
            for idx, item_num in enumerate(range(start, end + 1)):
                item = ITEMS[item_num]
                pre_score = pre_ratings.get(item_num, 0)
                post_score = post_ratings.get(item_num, 0)
                item_change = post_score - pre_score
                change_str = f"+{item_change}" if item_change > 0 else str(item_change)
                
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                    self._create_comparison_bar_chart(pre_score, post_score, colour_hex, tmp.name)
                    self._add_table_row(
                        items_table,
                        [str(item_num), item['text'], item['focus'], str(pre_score), str(post_score), None, change_str],
                        idx,
                        [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT,
                         WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                         WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                         WD_ALIGN_PARAGRAPH.CENTER],
                        bar_image_path=tmp.name,
                        bar_col=5,
                        col_widths=COL_WIDTHS_7
                    )
            
            # Reflection prompt after each indicator table
            doc.add_paragraph()
            reflect_table = doc.add_table(rows=1, cols=1)
            reflect_table.style = 'Table Grid'
            reflect_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            # Match width to the indicator table above (9000 twips)
            tbl = reflect_table._tbl
            tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
            tblLayout = OxmlElement('w:tblLayout')
            tblLayout.set(qn('w:type'), 'fixed')
            tblPr.append(tblLayout)
            # Remove any default indent
            tblInd = OxmlElement('w:tblInd')
            tblInd.set(qn('w:w'), '0')
            tblInd.set(qn('w:type'), 'dxa')
            tblPr.append(tblInd)
            reflect_cell = reflect_table.rows[0].cells[0]
            reflect_cell.width = 9000
            tc = reflect_cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), '9000')
            tcW.set(qn('w:type'), 'dxa')
            tcPr.append(tcW)
            self._set_cell_shading(reflect_cell, 'F5F5F5')
            self._set_cell_margins(reflect_cell, 100, 100, 100, 100)
            para = reflect_cell.paragraphs[0]
            run = para.add_run("Reflect: ")
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = COLOURS_RGB['purple']
            run.font.name = 'Arial'
            run = para.add_run("Looking at the changes above, what has shifted most for you in this area? "
                              "What will you continue to develop?")
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = COLOURS_RGB['mid_grey']
            run.font.name = 'Arial'
        
        # Page break before reflections
        doc.add_page_break()
        
        # Reflections
        heading = doc.add_paragraph()
        run = heading.add_run("Your Reflections")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = COLOURS_RGB['purple']
        
        intro = doc.add_paragraph()
        run = intro.add_run("Your responses to the open questions after the programme:")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = COLOURS_RGB['mid_grey']
        
        for q_num, question in OPEN_QUESTIONS_POST.items():
            # Question in a shaded box
            q_table = doc.add_table(rows=1, cols=1)
            q_table.style = 'Table Grid'
            q_cell = q_table.rows[0].cells[0]
            q_cell.text = question
            self._set_cell_shading(q_cell, 'F5F5F5')
            self._set_cell_margins(q_cell, 80, 80, 120, 120)
            for para in q_cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
                    run.font.color.rgb = COLOURS_RGB['purple']
            
            # For Q3, show original concern
            if q_num == 3:
                original = pre_responses.get(3, "")
                if original:
                    orig_para = doc.add_paragraph()
                    run = orig_para.add_run("Your original concern: ")
                    run.font.color.rgb = COLOURS_RGB['mid_grey']
                    run.font.size = Pt(9)
                    run = orig_para.add_run(f'"{original}"')
                    run.italic = True
                    run.font.size = Pt(9)
            
            response = post_responses.get(q_num, "No response provided")
            r_para = doc.add_paragraph()
            run = r_para.add_run(response)
            run.italic = True
            run.font.size = Pt(10)
            run.font.name = 'Arial'
            r_para.paragraph_format.space_after = Pt(12)
        
        # Closing
        closing = doc.add_paragraph()
        run = closing.add_run("For questions about your development plan, speak with your facilitator or line manager.")
        run.italic = True
        run.font.color.rgb = COLOURS_RGB['mid_grey']
        run.font.size = Pt(9)
        
        # Appendix
        self._add_appendix(doc)
        
        # Save
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    
    # =========== IMPACT REPORT ===========
    
    def generate_impact_report(self, cohort_id: int) -> io.BytesIO:
        """Generate an Impact report (cohort summary) with integrated AI insights."""
        
        # Get cohort data
        cohort_data = self.db.get_cohort_data(cohort_id)
        if not cohort_data:
            raise ValueError("Cohort not found")
        
        cohort = cohort_data['cohort']
        participants = cohort_data['participants']
        
        complete_participants = [p for p in participants if p['pre'] and p['post']]
        
        if len(complete_participants) < 2:
            raise ValueError("Need at least 2 participants with complete data")
        
        n_complete = len(complete_participants)
        
        # Calculate cohort averages
        pre_avgs = self.db.get_cohort_averages(cohort_id, 'PRE')
        post_avgs = self.db.get_cohort_averages(cohort_id, 'POST')
        
        pre_indicator_scores = {}
        post_indicator_scores = {}
        for indicator, (start, end) in INDICATORS.items():
            pre_scores = [pre_avgs.get(i, {}).get('avg', 0) for i in range(start, end + 1)]
            post_scores = [post_avgs.get(i, {}).get('avg', 0) for i in range(start, end + 1)]
            pre_indicator_scores[indicator] = sum(pre_scores) / len(pre_scores) if pre_scores else 0
            post_indicator_scores[indicator] = sum(post_scores) / len(post_scores) if post_scores else 0
        
        pre_overall = sum(pre_avgs.get(i, {}).get('avg', 0) for i in range(1, 33)) / 32
        post_overall = sum(post_avgs.get(i, {}).get('avg', 0) for i in range(1, 33)) / 32
        
        # Focus area analysis
        pre_focus = {}
        post_focus = {}
        for focus in FOCUS_TAGS.keys():
            items = get_items_by_focus(focus)
            pre_focus[focus] = sum(pre_avgs.get(i, {}).get('avg', 0) for i in items) / len(items) if items else 0
            post_focus[focus] = sum(post_avgs.get(i, {}).get('avg', 0) for i in items) / len(items) if items else 0
        
        # ===== DYNAMIC METRICS CALCULATION =====
        
        # Calculate % of participants who improved overall
        improved_count = 0
        for p in complete_participants:
            pre_ratings = p['pre']['ratings']
            post_ratings = p['post']['ratings']
            pre_avg = sum(pre_ratings.get(i, 0) for i in range(1, 33)) / 32
            post_avg = sum(post_ratings.get(i, 0) for i in range(1, 33)) / 32
            if post_avg > pre_avg:
                improved_count += 1
        pct_improved = (improved_count / n_complete * 100) if n_complete > 0 else 0
        
        # Calculate % of post-programme item averages at "Agree" (5) or above
        agree_count = 0
        total_items = 0
        for i in range(1, 33):
            post_avg = post_avgs.get(i, {}).get('avg', 0)
            if post_avg > 0:
                total_items += 1
                if post_avg >= 5.0:
                    agree_count += 1
        pct_agree_or_above = (agree_count / total_items * 100) if total_items > 0 else 0
        
        # Calculate item-level changes for top growth and lowest post
        item_changes = []
        for item_num in range(1, 33):
            pre_avg = pre_avgs.get(item_num, {}).get('avg', 0)
            post_avg = post_avgs.get(item_num, {}).get('avg', 0)
            item_changes.append({
                'num': item_num,
                'text': ITEMS[item_num]['text'],
                'focus': ITEMS[item_num]['focus'],
                'pre_avg': pre_avg,
                'post_avg': post_avg,
                'change': post_avg - pre_avg
            })
        
        top_growth_items = sorted(item_changes, key=lambda x: x['change'], reverse=True)[:5]
        lowest_post_items = sorted(item_changes, key=lambda x: x['post_avg'])[:5]
        
        completion_rate = int(cohort_data['post_completed'] / len(participants) * 100) if participants else 0
        
        # ===== COLLECT OPEN RESPONSES =====
        
        takeaway_responses = [p['post']['open_responses'].get(1, '') for p in complete_participants if p['post']['open_responses'].get(1)]
        commitment_responses = [p['post']['open_responses'].get(2, '') for p in complete_participants if p['post']['open_responses'].get(2)]
        concern_pre_responses = [p['pre']['open_responses'].get(3, '') for p in complete_participants if p['pre']['open_responses'].get(3)]
        concern_post_responses = [p['post']['open_responses'].get(3, '') for p in complete_participants if p['post']['open_responses'].get(3)]
        
        # ===== INTEGRATED AI INSIGHTS =====
        
        # Build the comprehensive score data package
        score_data = {
            'n_participants': n_complete,
            'pre_overall': pre_overall,
            'post_overall': post_overall,
            'indicator_scores': [
                {'name': ind, 'pre': pre_indicator_scores[ind], 'post': post_indicator_scores[ind],
                 'change': post_indicator_scores[ind] - pre_indicator_scores[ind]}
                for ind in INDICATORS.keys()
            ],
            'focus_scores': [
                {'name': foc, 'pre': pre_focus[foc], 'post': post_focus[foc],
                 'change': post_focus[foc] - pre_focus[foc]}
                for foc in FOCUS_TAGS.keys()
            ],
            'top_growth_items': top_growth_items,
            'lowest_post_items': lowest_post_items,
            'pct_improved': pct_improved,
            'pct_agree_or_above': pct_agree_or_above,
        }
        
        open_responses = {
            'takeaways': takeaway_responses,
            'commitments': commitment_responses,
            'concerns_pre': concern_pre_responses,
            'concerns_post': concern_post_responses,
        }
        
        # Single integrated API call for all insights
        insights = self.theme_extractor.extract_cohort_insights(score_data, open_responses)
        
        # ===== BUILD DOCUMENT =====
        
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(10)
        
        self._add_logo_header(doc)
        
        # Title
        title = doc.add_paragraph()
        run = title.add_run("THE READINESS FRAMEWORK")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = COLOURS_RGB['purple']
        
        subtitle = doc.add_paragraph()
        run = subtitle.add_run("Readiness Impact")
        run.font.size = Pt(12)
        run.font.color.rgb = COLOURS_RGB['magenta']
        
        # Cohort info
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        info_data = [
            ("Programme:", cohort.get('programme', 'Launch Readiness')),
            ("Cohort:", cohort['name']),
            ("Participants:", f"{len(participants)} enrolled | {cohort_data['pre_completed']} pre | {cohort_data['post_completed']} post"),
            ("Period:", f"{cohort.get('start_date', 'TBC')} to {cohort.get('end_date', 'TBC')}")
        ]
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = value
            self._set_cell_shading(info_table.rows[i].cells[0], 'F5F5F5')
            self._set_cell_margins(info_table.rows[i].cells[0])
            self._set_cell_margins(info_table.rows[i].cells[1])
            for cell in info_table.rows[i].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)
                        if cell == info_table.rows[i].cells[0]:
                            run.bold = True
        
        # Executive Summary — AI-generated narrative
        doc.add_paragraph()
        heading = doc.add_paragraph()
        run = heading.add_run("Executive Summary")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = COLOURS_RGB['purple']
        
        # Use AI narrative, splitting on double newlines for paragraphs
        executive_text = insights.get('executive_narrative', '')
        if executive_text:
            for para_text in executive_text.split('\n\n'):
                if para_text.strip():
                    p = doc.add_paragraph()
                    p.add_run(para_text.strip())
        else:
            change = post_overall - pre_overall
            p = doc.add_paragraph()
            p.add_run(f"The Launch Readiness programme delivered measurable improvement across all four "
                      f"Readiness Indicators. Cohort average scores increased from {pre_overall:.1f} to "
                      f"{post_overall:.1f} (+{change:.1f} on a 6-point scale).")
        
        # Key metrics boxes — NOW FULLY DYNAMIC
        change = post_overall - pre_overall
        doc.add_paragraph()
        metrics_table = doc.add_table(rows=1, cols=4)
        metrics_table.style = 'Table Grid'
        
        metrics = [
            (f"+{change:.1f}", "Average Increase", '461E96'),
            (f"{completion_rate}%", "Completion Rate", '00B4E6'),
            (f"{pct_improved:.0f}%", "Showed Improvement", 'E6008C'),
            (f"{pct_agree_or_above:.0f}%", "Now 'Agree' or Above", '00DC8C')
        ]
        
        for i, (value, label, colour) in enumerate(metrics):
            cell = metrics_table.rows[0].cells[i]
            self._set_cell_shading(cell, colour)
            self._set_cell_margins(cell, 80, 80, 60, 60)
            
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(value)
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = COLOURS_RGB['white']
            
            para2 = cell.add_paragraph()
            para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run2 = para2.add_run(label)
            run2.font.size = Pt(8)
            run2.font.color.rgb = COLOURS_RGB['white']
        
        # Page break
        doc.add_page_break()
        
        # Indicator Results
        heading = doc.add_paragraph()
        run = heading.add_run("Indicator Results")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = COLOURS_RGB['purple']
        
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            self._create_comparison_radar_chart(pre_indicator_scores, post_indicator_scores, tmp.name)
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()
            run.add_picture(tmp.name, width=Inches(4.5))
        
        doc.add_paragraph()
        results_table = self._create_styled_table(doc, ["Indicator", "Pre", "Post", "Change"])
        
        for i, indicator in enumerate(INDICATORS.keys()):
            pre = pre_indicator_scores.get(indicator, 0)
            post = post_indicator_scores.get(indicator, 0)
            change = post - pre
            change_str = f"+{change:.1f}" if change > 0 else f"{change:.1f}"
            self._add_table_row(results_table, [indicator, f"{pre:.1f}", f"{post:.1f}", change_str], i,
                               [WD_ALIGN_PARAGRAPH.LEFT] + [WD_ALIGN_PARAGRAPH.CENTER] * 3)
        
        # Focus Area Analysis
        doc.add_paragraph()
        heading = doc.add_paragraph()
        run = heading.add_run("Impact by Focus Area")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = COLOURS_RGB['purple']
        
        focus_intro = doc.add_paragraph()
        focus_intro.add_run("Each statement measures one of four focus areas. This shows where the programme had greatest impact.")
        
        doc.add_paragraph()
        focus_table = self._create_styled_table(doc, ["Focus Area", "What It Measures", "Pre", "Post", "Change"])
        
        focus_data = [
            ("Knowledge", "Understanding of concepts, processes, frameworks"),
            ("Awareness", "Recognition of own patterns, triggers, impact"),
            ("Confidence", "Self-belief and comfort in capability"),
            ("Behaviour", "Actions, habits and practices"),
        ]
        
        for i, (focus, desc) in enumerate(focus_data):
            pre = pre_focus.get(focus, 0)
            post = post_focus.get(focus, 0)
            change = post - pre
            change_str = f"+{change:.1f}" if change > 0 else f"{change:.1f}"
            self._add_table_row(focus_table, [focus, desc, f"{pre:.1f}", f"{post:.1f}", change_str], i,
                               [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT,
                                WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER,
                                WD_ALIGN_PARAGRAPH.CENTER])
        
        # Qualitative Themes — from integrated insights
        doc.add_paragraph()
        heading = doc.add_paragraph()
        run = heading.add_run("Qualitative Themes")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = COLOURS_RGB['purple']
        
        subheading = doc.add_paragraph()
        run = subheading.add_run("Most Valuable Takeaways")
        run.bold = True
        run.font.size = Pt(11)
        
        takeaway_themes = insights.get('takeaway_themes', [])
        themes = format_insight_themes(takeaway_themes, n_complete)
        if themes:
            for theme in themes:
                doc.add_paragraph(f"* {theme}")
        else:
            doc.add_paragraph("Manual review of responses recommended.")
        
        doc.add_paragraph()
        subheading = doc.add_paragraph()
        run = subheading.add_run("Commitments to Action")
        run.bold = True
        run.font.size = Pt(11)
        
        commitment_themes = insights.get('commitment_themes', [])
        themes = format_insight_themes(commitment_themes, n_complete)
        if themes:
            for theme in themes:
                doc.add_paragraph(f"* {theme}")
        else:
            doc.add_paragraph("Manual review of responses recommended.")
        
        # ROI Summary — AI-generated narrative
        doc.add_paragraph()
        heading = doc.add_paragraph()
        run = heading.add_run("ROI Summary")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = COLOURS_RGB['purple']
        
        roi_text = insights.get('roi_narrative', '')
        if not roi_text:
            # Fallback using actual data
            pre_q32 = pre_avgs.get(32, {}).get('avg', 0)
            post_q32 = post_avgs.get(32, {}).get('avg', 0)
            best_focus = max(score_data['focus_scores'], key=lambda x: x['change'])
            roi_text = (
                f"Before the programme, the average readiness score was {pre_overall:.1f}. "
                f"After completing Launch Readiness, this rose to {post_overall:.1f}. "
                f"{pct_agree_or_above:.0f}% of post-programme item scores now sit at 'Agree' or above. "
                f"The greatest gains were in {best_focus['name'].lower()} ({best_focus['change']:+.1f})."
            )
        
        roi_para = doc.add_paragraph()
        run = roi_para.add_run(roi_text)
        run.italic = True
        
        # Recommendations — AI-generated, data-driven
        doc.add_paragraph()
        subheading = doc.add_paragraph()
        run = subheading.add_run("Recommendations")
        run.bold = True
        run.font.size = Pt(11)
        
        recommendations = insights.get('recommendations', [])
        if not recommendations:
            # Fallback
            weakest = min(score_data['indicator_scores'], key=lambda x: x['post'])
            recommendations = [
                f"Reinforce {weakest['name']} — this indicator scored lowest post-programme ({weakest['post']:.1f})",
                "Ensure line managers support application of new frameworks in the workplace",
                "Consider 90-day follow-up assessment to measure sustained application",
                "Share anonymised cohort themes with participants to reinforce collective learning",
            ]
        
        for i, rec in enumerate(recommendations, 1):
            doc.add_paragraph(f"{i}. {rec}")
        
        # Closing
        doc.add_paragraph()
        closing = doc.add_paragraph()
        run = closing.add_run("For questions or to discuss follow-up interventions, contact the programme facilitators.")
        run.italic = True
        run.font.color.rgb = COLOURS_RGB['mid_grey']
        run.font.size = Pt(9)
        
        # Appendix
        self._add_appendix(doc)
        
        # Save
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
